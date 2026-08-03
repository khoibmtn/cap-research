#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tạo báo cáo phân tích và đề xuất nghiên cứu CAP
Output: cap_analysis_report.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── PAGE SETUP ───────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.0)

# ─── STYLES ───────────────────────────────────────────────────────────────────
styles = doc.styles

def set_style(style_name, font_name='Times New Roman', font_size=13,
              bold=False, color=None, space_before=0, space_after=6,
              line_spacing=None, alignment=None):
    try:
        style = styles[style_name]
    except:
        return
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor(*color)
    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)
    if alignment is not None:
        pf.alignment = alignment

set_style('Normal', font_size=13, space_after=6, line_spacing=20)

def add_heading(text, level=1, color=None):
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 70, 127)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(31, 73, 125)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        run.font.size = Pt(13)
        if color:
            run.font.color.rgb = RGBColor(*color)
        else:
            run.font.color.rgb = RGBColor(192, 80, 77)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(3)
    return p

def add_para(text, bold=False, italic=False, color=None, indent=0,
             size=13, space_after=4, align=None):
    p = doc.add_paragraph()
    p.style = 'Normal'
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(text, indent_level=1, bold_prefix=None):
    p = doc.add_paragraph()
    p.style = 'Normal'
    p.paragraph_format.left_indent = Cm(indent_level * 1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run_prefix = p.add_run(f'• {bold_prefix}')
        run_prefix.font.name = 'Times New Roman'
        run_prefix.font.size = Pt(13)
        run_prefix.font.bold = True
        run_rest = p.add_run(f' {text}')
        run_rest.font.name = 'Times New Roman'
        run_rest.font.size = Pt(13)
    else:
        run = p.add_run(f'• {text}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    hdr.height = Cm(0.9)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Background color
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        shade = 'EEF4FB' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), shade)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            is_bold = cell_text.startswith('**') and cell_text.endswith('**')
            txt = cell_text.strip('*') if is_bold else cell_text
            run = p.add_run(txt)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11.5)
            run.font.bold = is_bold

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_warning_box(text):
    """Tạo hộp cảnh báo (viền đỏ)"""
    p = doc.add_paragraph()
    p.style = 'Normal'
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top','left','bottom','right']:
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '12')
        bdr.set(qn('w:space'), '4')
        bdr.set(qn('w:color'), 'C0504D')
        pBdr.append(bdr)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 80, 77)
    return p

def add_info_box(text, fill='E8F4FD'):
    p = doc.add_paragraph()
    p.style = 'Normal'
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12.5)
    run.font.italic = True
    return p

def add_hr():
    p = doc.add_paragraph()
    p.style = 'Normal'
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4472C4')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
#  TRANG TIÊU ĐỀ
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('TRƯỜNG ĐẠI HỌC Y DƯỢC HẢI PHÒNG')
run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.font.bold = True
run.font.color.rgb = RGBColor(31, 73, 125)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
run = p.add_run('Nghiên cứu sinh: Nguyễn Thị Trang')
run.font.name = 'Times New Roman'; run.font.size = Pt(13)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BÁO CÁO PHÂN TÍCH NGHIÊN CỨU VÀ ĐỀ XUẤT')
run.font.name = 'Times New Roman'; run.font.size = Pt(18); run.font.bold = True
run.font.color.rgb = RGBColor(0, 70, 127)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Đặc điểm căn nguyên vi sinh và dấu ấn sinh học trong CAP')
run.font.name = 'Times New Roman'; run.font.size = Pt(15); run.font.bold = True
run.font.color.rgb = RGBColor(192, 80, 77)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Khoảng trống nghiên cứu • Phân tích bổ sung • Thu thập biến số mới')
run.font.name = 'Times New Roman'; run.font.size = Pt(13); run.font.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Tháng 6 năm 2026')
run.font.name = 'Times New Roman'; run.font.size = Pt(13)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  PHẦN I — TÓM TẮT ĐỀ TÀI
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('I. TÓM TẮT ĐỀ TÀI NGHIÊN CỨU', level=1)
add_hr()

add_heading('1.1. Thông tin đề tài', level=2)
add_table(
    headers=['Nội dung', 'Chi tiết'],
    rows=[
        ['Tên đề tài', 'Nghiên cứu đặc điểm căn nguyên vi sinh và một số dấu ấn sinh học ở bệnh nhân viêm phổi mắc phải cộng đồng (CAP) nhập viện'],
        ['Nghiên cứu sinh', 'Nguyễn Thị Trang'],
        ['Người hướng dẫn', 'PGS.TS Đào Ngọc Bằng; TS.BS Lê Thị Diệu Hiền'],
        ['Địa điểm', 'Bệnh viện Phổi Hải Phòng, 568 Trần Tất Văn, Phù Liễn, Hải Phòng'],
        ['Thời gian', '01/2026 – 12/2028'],
        ['Cỡ mẫu tối thiểu', 'n = 119 bệnh nhân (có hiệu chỉnh 5% từ chối)'],
        ['Thiết kế', 'Mô tả cắt ngang, có theo dõi dọc'],
        ['Tiêu chuẩn chọn mẫu', 'CAP người lớn ≥16 tuổi, PCR đa mồi dương tính với vi khuẩn'],
    ],
    col_widths=[5.5, 11]
)

doc.add_paragraph()
add_heading('1.2. Hai mục tiêu nghiên cứu', level=2)
add_bullet('Mô tả đặc điểm căn nguyên vi sinh và một số dấu ấn sinh học ở bệnh nhân CAP nhập viện điều trị tại Bệnh viện Phổi Hải Phòng.', bold_prefix='Mục tiêu 1:')
add_bullet('Đánh giá mối liên quan của các dấu ấn sinh học với mức độ nặng của viêm phổi mắc phải cộng đồng.', bold_prefix='Mục tiêu 2:')

doc.add_paragraph()
add_heading('1.3. Biomarker nghiên cứu', level=2)
add_table(
    headers=['Biomarker', 'Đơn vị', 'Phương pháp', 'Cơ sở xét nghiệm'],
    rows=[
        ['sTREM-1', 'pg/mL', 'ELISA', 'Labo Miễn dịch – Học viện Quân y'],
        ['TIMP-1', 'ng/mL', 'ELISA', 'Labo Miễn dịch – Học viện Quân y'],
        ['IL-6', 'pg/mL', 'ELISA', 'Labo Miễn dịch – Học viện Quân y'],
        ['IL-10', 'pg/mL', 'ELISA', 'Labo Miễn dịch – Học viện Quân y'],
        ['IL-17', 'pg/mL', 'ELISA', 'Labo Miễn dịch – Học viện Quân y'],
        ['CRP', 'mg/L', 'Sinh hóa tự động', 'Labo XN Bệnh viện Phổi Hải Phòng'],
        ['PCT', 'ng/mL', 'CLIA', 'Labo XN Bệnh viện Phổi Hải Phòng'],
        ['NLR', '—', 'Tính từ BCCTT/Lympho', 'Ứng dụng CAP Research'],
        ['PLR', '—', 'Tính từ Tiểu cầu/Lympho', 'Ứng dụng CAP Research'],
        ['CAR', '—', 'Tính từ CRP/Albumin', 'Ứng dụng CAP Research'],
    ],
    col_widths=[3.5, 2.5, 4, 6.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  PHẦN II — VẤN ĐỀ KHẨN CẤP
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('II. VẤN ĐỀ KHẨN CẤP: KHÔNG NHẤT QUÁN GIỮA ĐỀ CƯƠNG VÀ ỨNG DỤNG', level=1)
add_hr()

add_warning_box('⚠ CẢNH BÁO: Ứng dụng thu thập số liệu ĐANG THIẾU module PCR đa mồi — trong khi đây là TIÊU CHÍ CHỌN MẪU trong đề cương. Cần xử lý TRƯỚC KHI thu thập thêm bệnh nhân.')

doc.add_paragraph()
add_table(
    headers=['Vấn đề', 'Đề cương (tiêu chuẩn)', 'Ứng dụng (spec.md hiện tại)', 'Mức độ'],
    rows=[
        ['**PCR đa mồi**', 'Là TIÊU CHÍ LỰA CHỌN BN: PCR (+) với vi khuẩn mới vào NC', '"Biến số chưa có: PCR đa mồi (multiplex)"', '🔴 Khẩn cấp'],
        ['Kết quả vi sinh', 'PCR đa mồi + cấy vi khuẩn song song', 'Chỉ có cấy vi khuẩn + kháng sinh đồ', '🔴 Khẩn cấp'],
        ['PaO₂/FiO₂', 'Tiêu chuẩn ATS/IDSA nhập ICU (<250)', 'Không có FiO₂ → không tính được chỉ số', '🟡 Quan trọng'],
        ['APACHE II', 'Đề cập trong tổng quan và bảng dự kiến', 'Không có trong app', '🟡 Quan trọng'],
        ['CURB-65', '5 biến đã đủ trong app', 'Không tự tính, không hiển thị', '🟢 Cải tiến'],
    ],
    col_widths=[3.5, 5, 5, 2.5]
)

doc.add_paragraph()
add_heading('Giải pháp ưu tiên', level=3)
add_bullet('Bổ sung field PCR đa mồi: tên tác nhân, số lượng, phân loại (điển hình / không điển hình / virus), ngày xét nghiệm', bold_prefix='P1 – Ngay:')
add_bullet('Thêm field FiO₂ (%) tại lúc nhập viện để tính PaO₂/FiO₂ ratio tự động', bold_prefix='P2 – Ngay:')
add_bullet('Thêm nguồn bệnh phẩm vi sinh (đờm / dịch rửa phế quản / ngoáy tị hầu / máu)', bold_prefix='P3 – Ngay:')
add_bullet('Tính tự động CURB-65 từ 5 biến sẵn có (Glasgow, Ure, Nhịp thở, HA tâm thu, Tuổi)', bold_prefix='P4 – Cải tiến:')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  PHẦN III — TỔNG QUAN TÀI LIỆU
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('III. TỔNG QUAN TÀI LIỆU CÓ HỆ THỐNG', level=1)
add_hr()
add_info_box('Tổng hợp từ: Consensus MCP (session 1 + session 2), YHVN Vercel API, Web Search. Kiểm tra retraction qua WebSearch: Không có bài nào trong danh sách 30 tài liệu gốc bị retract. Cảnh báo: 2 bài sTREM-1 KHÔNG thuộc đề tài này bị Hindawi retract (2023–2024) do gian lận peer review — không ảnh hưởng đến tài liệu đang dùng. Ưu tiên năm 2017–2026.')

# 3.1 sTREM-1
add_heading('3.1. sTREM-1 trong CAP — Căn nguyên vi sinh và tiên lượng', level=2)
add_table(
    headers=['Tác giả (Năm)', 'Tạp chí', 'n', 'Kết quả chính', 'AUC/Cutoff'],
    rows=[
        ['Gibot et al. (2004)', 'N Engl J Med', '43', 'sTREM-1 trong BAL là marker độc lập mạnh nhất của viêm phổi. OR=41,5; Sensitivity 98%, Specificity 90%', 'AUC N/A; cutoff BAL >5 pg/mL'],
        ['Tejera et al. (2007)', 'Cytokine', '226', 'sTREM-1 là yếu tố tiên lượng độc lập tử vong nội viện CAP, sau khi đã điều chỉnh PSI, CURB-65, IL-6, IGF-1', 'AUROC 0.83'],
        ['How CK et al. (2011)', 'Am J Emerg Med', '120', 'sTREM-1 phân biệt CAP điển hình vs không điển hình. Median 65,2 vs 25,9 pg/mL (p<0,001)', 'AUC 0.87; cutoff 44,2 pg/mL; Se 81%, Sp 79%'],
        ['Wang Y et al. (2019)', 'Arch Gerontol Geriatr', 'N/A', 'sTREM-1 dự đoán tiên lượng nặng CAP người già tốt hơn CRP và điểm CPIS', 'AUC >CRP và CPIS'],
        ['Hogendoorn et al. (2022)', 'BMC Infect Dis', '110', 'PCT, IL-6, sTREM-1 đều có AUC xuất sắc phân biệt CAP vi khuẩn. Kết hợp Nhịp thở+PCT: Se 94%, Sp 82%', 'PCT 0.88; IL-6 0.84; sTREM-1 0.83'],
        ['Aladakatti et al. (2025)', 'Respir Med', '80 (VAP)', 'sTREM-1 đơn độc AUC=0,800; kết hợp với APACHE II → AUC=0,945. Cutoff >9,5 pg/mL', 'AUC 0.945 (tổ hợp)'],
        ['Esposito S et al. (2016)', 'PLoS ONE', '150 CAP', 'sTREM-1 phân biệt căn nguyên và đánh giá mức độ nặng CAP. Kết hợp với midregional proadrenomedullin (MR-proADM) tăng độ chính xác phân loại PSI. 40 lần trích dẫn (Scopus)', 'Phân biệt căn nguyên; kết hợp MR-proADM'],
        ['Mou S et al. (2022)', 'Can Respir J', 'CAP+COPD', 'Nồng độ sTREM-1 huyết thanh tăng có ý nghĩa ở CAP đồng mắc COPD. Mức sTREM-1 cao tương quan với mức độ tắc nghẽn và nặng của bệnh', 'sTREM-1 trong CAP+COPD'],
    ],
    col_widths=[3.5, 3.5, 1.2, 7, 3.3]
)

doc.add_paragraph()

# 3.2 TIMP-1
add_heading('3.2. TIMP-1 — Biomarker đặc thù giới tính trong tổn thương phổi cấp', level=2)
add_table(
    headers=['Tác giả (Năm)', 'Tạp chí', 'n', 'Kết quả chính', 'AUC/Đặc điểm'],
    rows=[
        ['Almuntashiri et al. (2023)', 'Chin Med J Pulm Crit Care Med', 'Review', 'TIMP-1 tăng cao trong CAP so với người khỏe; tương quan thuận với PSI, APACHE II. Dữ liệu hầu hết từ châu Âu/Bắc Mỹ — KHÔNG CÓ dữ liệu châu Á', 'Tổng quan hệ thống'],
        ['Almuntashiri et al. (2022)', 'Biol Sex Differ', '100 (ARDS)', 'TIMP-1 đặc hiệu cho NỮ trong ARDS/ALI. AUC tử vong 30 ngày ở nữ = 0,87; ở nam: không có tương quan có ý nghĩa', 'AUC 0.87 (nữ); cutoff 159,7 ng/mL'],
        ['Almuntashiri et al. (2024)', 'Physiol Rep', 'COVID-19, H1N1', 'TIMP-1 là gen nhạy cảm estrogen, nằm trên NST X. Nồng độ khác biệt nam/nữ trong COVID-19 và H1N1. Tương quan với PaO₂/FiO₂ chỉ ở nữ', 'Sex-specific, X-linked'],
        ['Jones et al. (2021)', 'Shock', '135 (Sepsis)', 'TIMP-1 liên quan độc lập với shock (OR 1,51), ARDS (OR 1,24), AKI (OR 1,18), tử vong (OR 1,20) trong sepsis', 'Per log increase'],
        ['Lorente et al. (2009)', 'Crit Care', 'Sepsis', 'TIMP-1 >531 ng/mL → tử vong tăng 80% (RR=1,80; 95% CI 1,13–2,87). Tương quan với SOFA, APACHE II, lactat', 'RR 1.80; cutoff 531 ng/mL'],
    ],
    col_widths=[3.5, 4, 1.5, 7, 3.5]
)
add_info_box('→ Khoảng trống: Gần như không có dữ liệu về TIMP-1 trong CAP người lớn tại châu Á / Việt Nam. Phân tích theo giới tính là điểm hoàn toàn mới.', fill='FFF3CD')

doc.add_paragraph()

# 3.3 IL-6, IL-10, IL-17
add_heading('3.3. IL-6, IL-10, IL-17 — Pattern cytokine theo căn nguyên vi sinh', level=2)
add_table(
    headers=['Tác giả (Năm)', 'Tạp chí', 'n', 'Pattern cytokine chính', 'Ứng dụng'],
    rows=[
        ['Menéndez et al. (2012)', 'Chest', '658', 'S. pneumoniae → PCT, IL-6 cao nhất. Vi khuẩn KĐH → PCT, IL-6 thấp hơn. Virus → PCT thấp, IL-10 cao. Legionella → CRP, TNF-α cao. Enterobacteriaceae → IL-8 cao', 'Phân biệt căn nguyên bằng cytokine profile'],
        ['Feng CM et al. (2021)', 'BMC Pulm Med', '186', 'IL-17 tương quan tất cả thang điểm nặng CAP. IL-17 cutoff tử vong = 86,80 ng/mL (AUC 0,89). IL-17 là yếu tố độc lập tiên lượng CAP nặng ở người lớn', 'Tiên lượng tử vong, nhập ICU'],
        ['Dao et al. (2023)', 'Pneumon', 'VN', 'IL-6 và IL-10 liên quan mức độ nặng CAP vi khuẩn tại Việt Nam. IL-10 cao hơn ở Gram dương. IL-6 giảm sau 7 ngày điều trị hiệu quả', 'Theo dõi đáp ứng điều trị'],
        ['Miyazaki et al. (2023)', 'Pneumonia', '210', 'Tỷ số AAT/IL-10 >65 dự đoán CAP vi khuẩn (OR=19,8; 95% CI 4,7–83,2). AUC panel 3 biến = 0,927', 'Phân biệt vi khuẩn vs virus'],
        ['Siljan et al. (2017)', 'Eur J Clin Invest', '247', 'IL-6 và MIP-1β là yếu tố tiên lượng độc lập kết quả xấu. IL-10 và IL-17 không độc lập sau điều chỉnh. Cytokine KHÔNG khác biệt theo căn nguyên trong NC này', 'Không phân biệt được căn nguyên'],
        ['Moravec et al. (2024)', 'Epidemiol Mikrobiol Imunol', '74 (sCAP)', 'IL-17A và Th17 theo 3 nhóm: VK, virus, đồng nhiễm. Không khác biệt về IL-17A giữa nhóm — nhưng Th17 máu ngoại vi cao tương quan với nguy cơ tử vong sCAP tăng. Nam giới và tuổi cao là yếu tố nguy cơ độc lập', 'Th17 máu dự báo tử vong sCAP'],
        ['Duan Y et al. (2025)', 'Ital J Pediatr', 'Trẻ em', 'M. pneumoniae → IL-6, IL-17A, IFN-γ cao hơn RSV và Adenovirus. IL-10 thấp hơn Adenovirus', 'Phân biệt M. pneumoniae vs virus'],
    ],
    col_widths=[3.3, 3.2, 1.2, 7, 3.8]
)

doc.add_paragraph()

# 3.4 NLR PLR CAR
add_heading('3.4. NLR, PLR, CAR — Chỉ số viêm tính toán từ xét nghiệm thường quy', level=2)
add_table(
    headers=['Chỉ số', 'Nghiên cứu (Năm)', 'n', 'Kết quả chính', 'AUC/Cutoff'],
    rows=[
        ['NLR', 'Cataudella et al. (2017)', '195 (cao tuổi)', 'NLR > 11,12 dự đoán tử vong 30 ngày tốt hơn PSI (p<0,05), CURB-65, CRP, WBC (p<0,001). Tử vong 30%: NLR 11,12–13,4; 50%: NLR 13,4–28,3; 100%: NLR >28,3. Gợi ý phân tầng nhập viện theo ngưỡng NLR cụ thể', 'AUC > PSI; 3 cutoff tầng nguy cơ'],
        ['NLR', 'Ganaie et al. (2025) – Meta-analysis', '~17.838', 'Pooled RR tử vong = 2,02 (95% CI 1,18–3,47). RR nhập ICU = 1,30 (1,11–1,53). Cutoff phổ biến nhất: NLR >10', 'Pooled RR 2.02'],
        ['NLR', 'Tekin et al. (2024)', '4.039', 'NLR >12 liên quan ICU (OR 1,405) nhưng KHÔNG cải thiện dự đoán tử vong sau khi đã có PSI. AUC PSI+NLR = 0,76 vs PSI = 0,75 (không có ý nghĩa LS)', 'AUC 0.76 (tổ hợp)'],
        ['NLR', 'Huang et al. (2025)', '812', 'Người cao tuổi CAP: cutoff NLR = 6,5 (tử vong 30 ngày). Thêm NLR vào SMARTCOP → AUC = 0,847. Tương quan phi tuyến, tăng tốc đến NLR=16', 'AUC 0.847 (SMARTCOP+NLR)'],
        ['PLR', 'Enersen et al. (2023)', '3.826', 'PLR cao liên quan độc lập tử vong 90 ngày (HR=1,001; 95% CI 1,000–1,001; p=0,029)', 'HR 1.001'],
        ['CAR', 'Luo et al. (2021)', 'CAP nặng', 'CAR độc lập dự đoán CAP nặng (OR=8,789; 95% CI 1,543–50,064). CAR trung bình: chứng 0,04; CAP nhẹ 0,68; CAP nặng-TB 1,37 mg/g', 'AUC 0.773; OR 8.8'],
        ['CAR', 'Ustaalioğlu et al. (2025)', '312', 'AUC của CAR cho tử vong 30 ngày = 0,837. Cutoff CAR >0,77: Se 75,9%, Sp 86,7%', 'AUC 0.837; cutoff 0.77'],
        ['CAR', 'Wang et al. (2025)', '117', 'Mô hình ML (MLP): CAR + tuổi + bạch cầu TT → AUC=0,89 dự đoán CAP tiến triển nặng. CAR đóng góp quan trọng nhất trong model', 'AUC 0.89 (ML model)'],
        ['PSI vs CAR', 'Ekşioğlu et al. (2025)', '349', 'PSI AUC = 0,884; CAR AUC = 0,677 (thua rõ). PSI vẫn vượt trội CAR đơn lẻ ở người cao tuổi', 'PSI 0.884 > CAR 0.677'],
        ['NLR/CAR/SII/CALLY', 'Baran et al. (2025)', '207', 'NLR (p=0,009), CAR (p=0,011), CLR (p=0,006), SII (p=0,013) cao hơn ở không sống sót. CALLY index (CRP × Albumin × Lympho) có giá trị dự đoán tử vong tốt nhất trong panel (p=0,003). Tử vong 11%, lưu viện trung vị 8 ngày', 'CALLY > NLR > CAR; mới nhất'],
        ['NLR (Mycoplasma)', 'Jiang et al. (2026) – Systematic review', 'Trẻ em CAP', 'Systematic review NLR trong Mycoplasma pneumoniae pneumonia ở trẻ em. NLR tương quan với mức độ nặng, đáp ứng điều trị. NLR ngưỡng thấp hơn ở trẻ em so với người lớn', 'NLR trẻ em M. pneumoniae'],
    ],
    col_widths=[1.5, 3.5, 1.5, 7.5, 2.5]
)

doc.add_paragraph()

# 3.5 Vietnam
add_heading('3.5. Căn nguyên vi sinh CAP tại Việt Nam và Đông Nam Á', level=2)
add_table(
    headers=['Tác giả (Năm)', 'Cơ sở/Vùng', 'n', 'Căn nguyên phổ biến nhất', 'Kháng thuốc'],
    rows=[
        ['Phan HT Vy et al. (2025)', 'BV ĐK TW Cần Thơ', 'CAP nặng', 'PCR đa mồi phát hiện 92,2% tác nhân; đồng nhiễm 76,5%. RSV là virus phổ biến nhất', 'PCR vượt trội nuôi cấy'],
        ['Lý Khánh Vân et al. (2025)', 'Việt Nam', 'CAP nhập viện', 'K. pneumoniae, A. baumannii, H. influenzae, E. coli, S. pneumoniae (Top 5)', 'MDR cao'],
        ['Tran et al. (2022)', 'Vĩnh Long', '254', 'Cấy đàm (+) 61,8%. S. pneumoniae 12,6%; K. pneumoniae 12,2%; P. aeruginosa 8,3%. Enterobacteriaceae 36,5%', 'Kháng macrolide của Phế cầu 84,4%'],
        ['Tran TN Dung et al. (2025)', '6 BV lớn – Toàn quốc', '1.000 isolates', 'Chủ yếu HAP: A. baumannii 49,6%; P. aeruginosa 21%; K. pneumoniae 18,6%', 'MDR: A. baumannii 96%; K. pneumoniae 78%; kháng Ceftazidime-avibactam 34,3%'],
        ['Nguyễn TH et al. (2025)', 'BV Bạch Mai', 'CAP nhập viện', 'Realtime PCR đa mồi: 39,2% dương tính virus. RSV chiếm ưu thế. Đồng nhiễm Influenza+RSV 9,68%', 'PCR phát hiện tốt hơn cấy'],
        ['Bạch Thái Dương et al. (2024)', 'BV ĐK TW Cần Thơ\n(Tạp chí YHVN)', 'Người lớn', 'Vi khuẩn đa kháng thuốc trong CAP người lớn nhập viện tại tuyến tỉnh Miền Nam. Phổ MDR và đề kháng theo tác nhân phân lập được', 'Dữ liệu MDR tuyến tỉnh VN'],
        ['Lê Thị Diệu Hiền et al. (2021)', 'BV Phổi / BVQY\n(Tạp chí YHVN)', 'CAP vi khuẩn VN', 'Thay đổi nồng độ cytokine huyết thanh ở bệnh nhân CAP vi khuẩn. Dữ liệu tại Việt Nam về IL-6, IL-10, TNF-α trong CAP vi khuẩn', 'Dữ liệu cytokine CAP VN'],
        ['Đỗ Thanh Hoà & Lê Đức Giang (2025)', 'Tạp chí YHVN\n2025;557(1)', 'CAP nhập viện', 'So sánh giá trị một số thang điểm (PSI, CURB-65, SMART-COP) trong dự đoán nhập đơn vị hồi sức tích cực ở bệnh nhân CAP — bối cảnh Việt Nam', 'PSI vs CURB-65 tại VN'],
        ['Lê Thị Huệ et al. (2023)', 'Tạp chí YHVN\n526(2)', 'CAP nhập viện', 'Multiplex PCR phát hiện đồng nhiễm vi khuẩn + virus ở bệnh nhân CAP nhập viện. Tỷ lệ đồng nhiễm và phân bố tác nhân theo phương pháp PCR đa mồi', 'Dữ liệu đồng nhiễm VN'],
        ['Lý Khánh Vân et al. (2023)', 'Tạp chí YHVN\n530(1)', 'CAP nhập viện', 'Tác nhân vi sinh phát hiện trên mẫu đàm bằng PCR đa mồi. K. pneumoniae, A. baumannii, H. influenzae chiếm đa số', 'Phân bố tác nhân PCR đàm'],
        ['Lý Khánh Vân & Phạm Hùng Vân (2025)', 'Tạp chí YHVN\n548(1)', 'CAP nhập viện', 'Carbapenemase genes và kháng kháng sinh A. baumannii trong CAP nhập viện tại Việt Nam — gene NDM, OXA-23 phổ biến', 'MDR A. baumannii gene carbapenemase'],
        ['Nguyễn Thị Pháp & Phan Vũ Nguyên (2024)', 'Tạp chí YHVN\n545(1)', '145 CAP', 'So sánh ATS/IDSA vs CURB-65 vs PSI tiên lượng tử vong 30 ngày tại BV Phạm Ngọc Thạch. ATS/IDSA: AUC=0,902 (độ nhạy 92,9%), CURB-65: AUC=0,89, PSI: AUC=0,85', 'ATS/IDSA > PSI/CURB-65 tại VN'],
        ['Goyet et al. (2014)', 'VN, Thailand, Cambodia', '5.919', 'S. pneumoniae và H. influenzae phổ biến nhất. Tại VN: Mycoplasma và Chlamydia cũng thường gặp. B. pseudomallei quan trọng trong CAP nặng', 'Kháng penicillin phế cầu cao'],
    ],
    col_widths=[3.5, 3, 1.5, 6.5, 4]
)

# 3.6 MDR bacteria
add_heading('3.6. Vi khuẩn đa kháng thuốc (MDR) và đáp ứng viêm hệ thống', level=2)
add_para('Câu hỏi còn bỏ ngỏ: Vi khuẩn MDR có tạo ra phản ứng viêm hệ thống (biomarker) khác với VK nhạy cảm không? Và điều đó có dẫn đến kết cục xấu hơn không? Đây là nền tảng khoa học cho Bảng B4 đề xuất.')
add_table(
    headers=['Tác giả (Năm)', 'Tạp chí', 'n', 'So sánh MDR vs non-MDR', 'Kết quả quan trọng'],
    rows=[
        ['Karamouzos et al. (2021)', 'Infectious Diseases', '128 sepsis ICU', 'MDR (n=38) vs non-MDR (n=90) Gram âm sepsis. Đo: sTREM-1, IL-6, IL-10, TNF-α, Ang-2 khi nhập viện', 'TNF-α THẤP HƠN ở MDR (p=0,017). Carbapenem resistance: OR tử vong 5,38. Survival 28 ngày THẤP HƠN ở MDR (p=0,008)'],
        ['Wang Han et al. (2024)', 'eBioMedicine', 'Chuột + invitro', 'A. baumannii MDR — cơ chế cytokine storm. Đại thực bào phế nang và kẽ là đích tấn công chính', 'A. baumannii kích hoạt M1 macrophage qua TLR2/Myd88/NF-κB → cytokine storm → tử vong. Naproxen ức chế phân cực M1 → bảo vệ toàn phần. Gợi ý tiêu chí sinh học phân biệt nhiễm A. baumannii'],
        ['Kumar NR et al. (2024)', 'Antibiotics (Review)', '11 triệu ca/năm', 'Review toàn diện MDR sepsis: epidemiology, immune response, điều trị', 'MDR tạo immune dysregulation khác với VK nhạy cảm. Antibiotic overuse + underlying comorbidities → immune paralysis. POCT biomarker có thể phát hiện kháng thuốc tại giường bệnh'],
    ],
    col_widths=[3.5, 3, 2, 5.5, 3.5]
)
add_info_box('→ Khoảng trống tại Việt Nam: Mặc dù MDR rất phổ biến (A. baumannii 96%, K. pneumoniae 78%), chưa có nghiên cứu nào so sánh profile biomarker (sTREM-1, IL-6, NLR) giữa VK nhạy cảm và MDR trong CAP tại Việt Nam. Dữ liệu kháng sinh đồ đã có trong app → có thể phân tích ngay khi đủ cỡ mẫu.', fill='FFF3CD')

doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  PHẦN IV — KHOẢNG TRỐNG NGHIÊN CỨU
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('IV. KHOẢNG TRỐNG NGHIÊN CỨU VÀ GIÁ TRỊ ĐỀ TÀI', level=1)
add_hr()

# KT1
add_heading('Khoảng trống 1: Không có dữ liệu TIMP-1 trong CAP tại Việt Nam và châu Á', level=3)
add_para('Tổng quan hệ thống Almuntashiri 2023 ghi nhận TIMP-1 tăng cao trong CAP và tương quan với PSI, APACHE II, nhưng toàn bộ dữ liệu đến từ châu Âu và Bắc Mỹ. Không có nghiên cứu nào từ Việt Nam hay Đông Nam Á [TK 8,9].', size=13)
add_para('Đặc biệt quan trọng: TIMP-1 là gene nằm trên nhiễm sắc thể X, nhạy cảm với estrogen [TK 10]. Almuntashiri et al. 2022 (ARDS, n=100) cho thấy ở nữ, TIMP-1 có AUC 0,87 dự đoán tử vong 30 ngày (cutoff 159,7 ng/mL; Se 100%, Sp 74%), trong khi ở nam không có tương quan có ý nghĩa [TK 8]. Phân tích phân tầng theo giới này chưa từng được kiểm định trong bối cảnh CAP người lớn tại châu Á.', size=13)
add_info_box('→ Giá trị mới: Đề tài này sẽ tạo ra dữ liệu đầu tiên về TIMP-1 trong CAP tại Việt Nam và có tiềm năng xác nhận hiệu ứng giới tính trong bối cảnh khác với ARDS.', fill='E2EFDA')

# KT2
add_heading('Khoảng trống 2: Chưa có "pattern cytokine theo căn nguyên" tại Việt Nam', level=3)
add_para('Menéndez et al. (Chest 2012, n=658) là nghiên cứu nền tảng về pattern cytokine theo loại tác nhân trong CAP — nhưng được thực hiện tại Tây Ban Nha với phổ vi khuẩn hoàn toàn khác (S. pneumoniae chiếm ưu thế) [TK 4]. Tại Việt Nam, K. pneumoniae, A. baumannii, và H. influenzae phổ biến hơn [TK 29, 30]. Không có nghiên cứu nào kiểm định pattern IL-6/IL-10/IL-17/sTREM-1 theo từng loại tác nhân trong bối cảnh này.', size=13)
add_info_box('→ Giá trị mới: Phân tích pattern cytokine theo tác nhân PCR (điển hình, không điển hình, virus, đồng nhiễm) tại Việt Nam — điều này chưa ai thực hiện.', fill='E2EFDA')

# KT3
add_heading('Khoảng trống 3: Ngưỡng cắt NLR/PLR/CAR chưa xác định cho dân số Việt Nam', level=3)
add_para('Meta-analysis Ganaie 2025 (n=17.838 BN) xác nhận NLR >10 liên quan tử vong (RR 2,02) nhưng ngưỡng cắt thay đổi từ 4 đến 13,4 tùy quần thể [TK 18]. Nghiên cứu Huang 2025 ở người cao tuổi châu Á cho cutoff NLR = 6,5 — khác biệt so với cutoff >10 ở phương Tây [TK 20]. CAR có AUC 0,837 cho tử vong 30 ngày (Ustaalioğlu 2025) với cutoff 0,77 — chưa có dữ liệu tương đương từ Việt Nam [TK 22].', size=13)
add_info_box('→ Giá trị mới: Xác định ngưỡng cắt NLR, PLR, CAR tối ưu cho bệnh nhân CAP nhập viện tại Hải Phòng — ứng dụng lâm sàng trực tiếp, không tốn thêm chi phí xét nghiệm.', fill='E2EFDA')

# KT4
add_heading('Khoảng trống 4: Chưa có mô hình tổ hợp PSI + biomarker mới tại Việt Nam', level=3)
add_para('Menéndez 2009 (Thorax, 229 trích dẫn) cho thấy thêm CRP vào PSI cải thiện AUC từ 0,80 lên 0,85 [TK 28]. Çetin 2025 chứng minh PSI + FAR + CT severity score đạt AUC 0,844 cho tử vong ICU [TK 27]. Tổ hợp sTREM-1 + APACHE II nâng AUC lên 0,945 trong VAP [TK 1]. Không có mô hình nào tích hợp biomarker mới (sTREM-1, TIMP-1, IL-6) với PSI trong bối cảnh CAP tại Việt Nam.', size=13)
add_info_box('→ Giá trị mới: Xây dựng mô hình tổ hợp PSI + biomarker tối ưu cho phân tầng nguy cơ CAP tại Bệnh viện Phổi Hải Phòng — thực tế và khả thi.', fill='E2EFDA')

# KT5
add_heading('Khoảng trống 5: Mối liên quan kháng kháng sinh (MDR) — biomarker — kết cục', level=3)
add_para('Tỷ lệ MDR tại Việt Nam cực cao: A. baumannii 96%, K. pneumoniae 78%, P. aeruginosa 57% [TK 30]. Kháng ceftazidime-avibactam đã xuất hiện (K. pneumoniae 34,3%). Câu hỏi "VK kháng thuốc có làm tăng mức độ viêm hệ thống (biomarker cao hơn) và worsens kết cục không?" chưa có câu trả lời trong y văn Việt Nam.', size=13)
add_info_box('→ Giá trị mới: Phân tích biomarker theo nhóm VK nhạy cảm vs. MDR/XDR — gợi ý cơ chế và hỗ trợ quyết định kháng sinh sớm.', fill='E2EFDA')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  PHẦN V — ĐỀ XUẤT PHÂN TÍCH BỔ SUNG
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('V. ĐỀ XUẤT PHÂN TÍCH BỔ SUNG (thực hiện trên ứng dụng)', level=1)
add_hr()
add_para('Tất cả các bảng sau đây đều khả thi với biến số đang thu thập, không yêu cầu xét nghiệm bổ sung (ngoại trừ nhóm C cần PCR data).')

# Nhóm A
add_heading('Nhóm A — Bổ sung cho Mục tiêu 1 (Mô tả đặc điểm)', level=2)

add_heading('Bảng MỚI A1: Nồng độ biomarker theo giới tính [Khoảng trống 1]', level=3)
add_table(
    headers=['Chỉ số', 'Nam (n=?)', 'Nữ (n=?)', 'p-value'],
    rows=[
        ['sTREM-1 (pg/mL)', 'Median (Q1–Q3)', 'Median (Q1–Q3)', 'Mann-Whitney U'],
        ['**TIMP-1 (ng/mL)**', 'Median (Q1–Q3)', 'Median (Q1–Q3)', 'Mann-Whitney U'],
        ['IL-6 (pg/mL)', 'Median (Q1–Q3)', 'Median (Q1–Q3)', 'Mann-Whitney U'],
        ['IL-10 (pg/mL)', 'Median (Q1–Q3)', 'Median (Q1–Q3)', 'Mann-Whitney U'],
        ['IL-17 (pg/mL)', 'Median (Q1–Q3)', 'Median (Q1–Q3)', 'Mann-Whitney U'],
        ['CRP (mg/L)', 'Median (Q1–Q3)', 'Median (Q1–Q3)', 'Mann-Whitney U'],
        ['PCT (ng/mL)', 'Median (Q1–Q3)', 'Median (Q1–Q3)', 'Mann-Whitney U'],
        ['NLR', 'Median (Q1–Q3)', 'Median (Q1–Q3)', 'Mann-Whitney U'],
        ['CAR', 'Median (Q1–Q3)', 'Median (Q1–Q3)', 'Mann-Whitney U'],
    ],
    col_widths=[4, 4, 4, 4.5]
)
add_info_box('Ý nghĩa: TIMP-1 dự kiến cao hơn có ý nghĩa ở nữ (do gene trên NST X). Đây là phân tích hoàn toàn mới trong y văn Việt Nam.')

doc.add_paragraph()
add_heading('Bảng MỚI A2: So sánh biomarker theo loại căn nguyên PCR [Khoảng trống 2]', level=3)
add_table(
    headers=['Chỉ số', 'VK điển hình', 'VK không điển hình', 'Virus', 'Đồng nhiễm ≥2', 'p (Kruskal-Wallis)'],
    rows=[
        ['sTREM-1', 'Median (IQR)', 'Median (IQR)', 'Median (IQR)', 'Median (IQR)', ''],
        ['IL-6', 'Median (IQR)', 'Median (IQR)', 'Median (IQR)', 'Median (IQR)', ''],
        ['IL-10', 'Median (IQR)', 'Median (IQR)', 'Median (IQR)', 'Median (IQR)', ''],
        ['IL-17', 'Median (IQR)', 'Median (IQR)', 'Median (IQR)', 'Median (IQR)', ''],
        ['PCT', 'Median (IQR)', 'Median (IQR)', 'Median (IQR)', 'Median (IQR)', ''],
        ['CRP', 'Median (IQR)', 'Median (IQR)', 'Median (IQR)', 'Median (IQR)', ''],
    ],
    col_widths=[2.8, 2.5, 2.5, 2.5, 2.5, 2.7]
)
add_info_box('Điều kiện: Cần bổ sung PCR data vào app. Dùng Kruskal-Wallis + post-hoc Dunn. Pattern dự kiến: PCT và IL-6 cao nhất ở VK điển hình; IL-10 cao ở virus.')

doc.add_paragraph()
add_heading('Bảng MỚI A3: Tỷ lệ đồng nhiễm và liên quan mức độ nặng', level=3)
add_table(
    headers=['Chỉ tiêu', 'Đơn nhiễm (n=?)', 'Đồng nhiễm ≥2 (n=?)', 'p-value'],
    rows=[
        ['PSI IV–V (nặng)', 'n (%)', 'n (%)', 'Chi-square'],
        ['Tử vong nội viện', 'n (%)', 'n (%)', 'Fisher exact'],
        ['Thở máy', 'n (%)', 'n (%)', 'Fisher exact'],
        ['Sốc nhiễm khuẩn', 'n (%)', 'n (%)', 'Fisher exact'],
        ['Số ngày điều trị', 'Median (IQR)', 'Median (IQR)', 'Mann-Whitney U'],
        ['sTREM-1 (pg/mL)', 'Median (IQR)', 'Median (IQR)', 'Mann-Whitney U'],
    ],
    col_widths=[5, 3.5, 3.5, 4.5]
)

doc.add_paragraph()

# Nhóm B
add_heading('Nhóm B — Bổ sung cho Mục tiêu 2 (Liên quan biomarker – mức độ nặng)', level=2)

add_heading('Bảng MỚI B1: ROC so sánh các biomarker — 3 endpoint [Ứng dụng lâm sàng cao nhất]', level=3)
add_table(
    headers=['Chỉ số', 'AUC (95% CI)', 'Cutoff tối ưu', 'Độ nhạy (%)', 'Độ đặc hiệu (%)', 'p (DeLong vs PSI)'],
    rows=[
        ['**PSI Score (tham chiếu)**', '... (0.??–0.??)', '...', '...', '...', 'Ref'],
        ['NLR', '... (0.??–0.??)', '...', '...', '...', '...'],
        ['PLR', '... (0.??–0.??)', '...', '...', '...', '...'],
        ['CAR', '... (0.??–0.??)', '...', '...', '...', '...'],
        ['sTREM-1', '... (0.??–0.??)', '...', '...', '...', '...'],
        ['TIMP-1', '... (0.??–0.??)', '...', '...', '...', '...'],
        ['IL-6', '... (0.??–0.??)', '...', '...', '...', '...'],
        ['IL-17', '... (0.??–0.??)', '...', '...', '...', '...'],
        ['CRP', '... (0.??–0.??)', '...', '...', '...', '...'],
        ['PCT', '... (0.??–0.??)', '...', '...', '...', '...'],
    ],
    col_widths=[3, 3, 2.5, 2, 2, 4]
)
add_info_box('Thực hiện 3 bảng ROC riêng cho 3 endpoint: (A) PSI IV–V; (B) Tử vong nội viện; (C) Thở máy. So sánh AUC bằng DeLong test (đã có trong kế hoạch phân tích của đề cương).')

doc.add_paragraph()
add_heading('Bảng MỚI B2: Mô hình tổ hợp PSI + biomarker [Khoảng trống 4]', level=3)
add_table(
    headers=['Mô hình', 'AUC (95% CI)', 'ΔAUC so với PSI', 'p (DeLong)', 'AIC'],
    rows=[
        ['**PSI đơn thuần (tham chiếu)**', '0.??', 'Ref', '—', '...'],
        ['PSI + NLR', '0.??', '+0.0?', '...', '...'],
        ['PSI + CAR', '0.??', '+0.0?', '...', '...'],
        ['PSI + sTREM-1', '0.??', '+0.0?', '...', '...'],
        ['PSI + IL-6', '0.??', '+0.0?', '...', '...'],
        ['PSI + sTREM-1 + IL-6', '0.??', '+0.0?', '...', '...'],
        ['PSI + sTREM-1 + TIMP-1 + IL-6', '0.??', '+0.0?', '...', '...'],
    ],
    col_widths=[5, 2.5, 2.5, 2.5, 4]
)
add_info_box('Dùng hồi quy logistic đa biến trên SPSS. Endpoint: PSI IV–V hoặc tử vong. Mục tiêu: tìm tổ hợp tối ưu, có khả năng nâng AUC của PSI lên ≥0,85.')

doc.add_paragraph()
add_heading('Bảng MỚI B3: TIMP-1 phân tầng theo giới — ROC mortality [Khoảng trống 1 – Điểm mới nhất]', level=3)
add_table(
    headers=['', 'Toàn bộ BN', 'Nam', 'Nữ'],
    rows=[
        ['n', '...', '...', '...'],
        ['TIMP-1 Median (IQR), ng/mL', '...', '...', '...'],
        ['**AUC dự đoán tử vong 30 ngày (95% CI)**', '...', '...', '...'],
        ['Cutoff tối ưu (ng/mL)', '...', '...', '...'],
        ['Độ nhạy (%)', '...', '...', '...'],
        ['Độ đặc hiệu (%)', '...', '...', '...'],
        ['p (DeLong nam vs nữ)', '—', '—', '...'],
    ],
    col_widths=[5.5, 3.5, 3.5, 4]
)
add_info_box('Đây là bảng có tính mới NHẤT trong toàn bộ đề xuất. Nếu TIMP-1 ở nữ có AUC cao hơn đáng kể (như dữ liệu ARDS: 0,87 vs không có ở nam) → finding đủ mạnh để đăng tạp chí quốc tế (Q1 Respiratory Medicine / Critical Care).', fill='FFF3CD')

doc.add_paragraph()
add_heading('Bảng MỚI B4: Biomarker theo nhóm kháng sinh đồ [Khoảng trống 5]', level=3)
add_table(
    headers=['Chỉ số', 'VK nhạy cảm (n=?)', 'VK MDR/XDR (n=?)', 'p-value'],
    rows=[
        ['sTREM-1 (pg/mL)', 'Median (IQR)', 'Median (IQR)', 'Mann-Whitney U'],
        ['TIMP-1 (ng/mL)', 'Median (IQR)', 'Median (IQR)', ''],
        ['IL-6 (pg/mL)', 'Median (IQR)', 'Median (IQR)', ''],
        ['IL-10 (pg/mL)', 'Median (IQR)', 'Median (IQR)', ''],
        ['NLR', 'Median (IQR)', 'Median (IQR)', ''],
        ['CAR', 'Median (IQR)', 'Median (IQR)', ''],
        ['Tử vong nội viện', 'n (%)', 'n (%)', 'Chi-square'],
        ['Thở máy', 'n (%)', 'n (%)', 'Fisher'],
        ['Số ngày điều trị', 'Mean ± SD', 'Mean ± SD', 'Student t/MW-U'],
    ],
    col_widths=[4, 3.5, 3.5, 5.5]
)

doc.add_paragraph()
add_heading('Bảng MỚI C: Mở rộng tương quan Spearman (Bảng 3.11 hiện tại)', level=3)
add_table(
    headers=['Cặp tương quan', 'Spearman r', 'Phân loại mức độ', 'p-value'],
    rows=[
        ['Biomarker vs Số ngày lọc máu (nếu n đủ)', '', '', ''],
        ['sTREM-1 vs Số ngày từ khởi bệnh đến nhập viện', '', '', ''],
        ['NLR vs Số ngày điều trị', '', '', ''],
        ['CAR vs Số ngày điều trị', '', '', ''],
        ['IL-17 vs PSI Score', '', '', ''],
        ['TIMP-1 vs Số ngày thở máy (nếu có)', '', '', ''],
    ],
    col_widths=[7, 2.5, 3, 4]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  PHẦN VI — ĐỀ XUẤT THU THẬP THÊM
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('VI. ĐỀ XUẤT THU THẬP THÊM BIẾN SỐ MỚI', level=1)
add_hr()

add_heading('6.1. Ưu tiên 1 — Bắt buộc cho tính toàn vẹn đề tài', level=2)
add_table(
    headers=['Biến số', 'Mô tả', 'Phân tích được phép', 'Cách thu thập trong app'],
    rows=[
        ['**Kết quả PCR đa mồi**', 'Danh sách tác nhân, số lượng, ngày XN', 'Bảng A2, A3 — pattern cytokine theo VK; tỷ lệ đồng nhiễm', 'Multi-select: tên tác nhân (PCR); phân loại tự động: điển hình/KĐH/virus'],
        ['**FiO₂ tại nhập viện (%)**', 'Phân số oxy hít vào lúc nhập viện', 'Tính PaO₂/FiO₂ ratio (tiêu chuẩn ATS <250); phân loại suy hô hấp', 'Field số: FiO₂ (%), default = 21% nếu thở khí trời'],
        ['**Nguồn bệnh phẩm vi sinh**', 'Loại mẫu xét nghiệm vi sinh', 'Phân tích tỷ lệ dương tính theo nguồn; kiểm soát sai số chọn lọc', 'Dropdown: Đờm / Dịch rửa phế quản / Ngoáy tị hầu / Máu / Khác'],
    ],
    col_widths=[3.5, 3.5, 5, 4.5]
)

doc.add_paragraph()
add_heading('6.2. Ưu tiên 2 — Mở rộng phân tích', level=2)
add_table(
    headers=['Biến số', 'Mô tả', 'Phân tích mới được phép', 'Giá trị khoa học'],
    rows=[
        ['**CURB-65**', '5 biến đã ĐỦ trong app (Glasgow, Ure, Nhịp thở, HA tâm thu, Tuổi)', 'So sánh PSI vs CURB-65 vs Biomarker trong phân tầng nguy cơ; tính AUC mỗi thang điểm', 'Thực hành: CURB-65 đơn giản hơn PSI; BV tuyến tỉnh cần biết thang nào tốt hơn'],
        ['Loại oxy hỗ trợ', 'Thở khí trời / Kính mũi / Mặt nạ / HFNC / NIPPV / Thở máy', 'Phân loại suy hô hấp theo thang độ; tương quan với biomarker', 'Chưa có dữ liệu kiểu này từ Việt Nam'],
        ['Thời điểm lấy mẫu biomarker', 'Giờ thứ mấy sau nhập viện?', 'Chuẩn hóa comparison; giảm bias khi so sánh nhóm', 'Kiểm soát sai số đo lường'],
    ],
    col_widths=[3.5, 4, 4.5, 4.5]
)

doc.add_paragraph()
add_heading('6.3. Ưu tiên 3 — Phân tích dọc (Longitudinal) — Giá trị cao nhất, khó nhất', level=2)
add_table(
    headers=['Biến số', 'Thời điểm lấy mẫu', 'Phân tích', 'Giá trị khoa học'],
    rows=[
        ['**Biomarker lần 2** (sTREM-1, IL-6, TIMP-1)', 'Ngày 3–5 điều trị', 'Δ biomarker (giảm/tăng) → tiên lượng đáp ứng điều trị; Biomarker kinetics', 'Dao 2023 đã chứng minh IL-6 giảm sau 7 ngày ĐT hiệu quả. Chưa có dữ liệu dọc từ VN'],
        ['**Biomarker lần 3** (tùy chọn)', 'Khi ra viện', 'Ai trở về bình thường? Ai vẫn cao? Liên quan đến tái nhập viện?', 'Hoàn toàn mới tại Việt Nam; có thể thành bài báo riêng về biomarker kinetics'],
    ],
    col_widths=[4, 2.5, 5, 5]
)
add_warning_box('Lưu ý: Phân tích dọc yêu cầu (1) lấy máu thêm ngày 3–5 và khi ra viện, (2) cập nhật protocol nghiên cứu, (3) bổ sung field timestamp riêng trong app. Chi phí cao nhưng giá trị công bố cao nhất — có thể tách thành bài báo độc lập.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  PHẦN VII — BẢN ĐỒ TRIỂN KHAI
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('VII. BẢN ĐỒ ƯU TIÊN TRIỂN KHAI', level=1)
add_hr()

add_heading('7.1. Theo giai đoạn nghiên cứu', level=2)
add_table(
    headers=['Giai đoạn', 'Điều kiện', 'Việc cần làm', 'Phân tích được phép'],
    rows=[
        ['**Ngay lập tức**\n(Trước khi thu thập thêm BN)', 'Ứng dụng', 'Bổ sung module PCR đa mồi, FiO₂, nguồn BN, tự tính CURB-65', 'Không có phân tích mới (chỉ sửa app)'],
        ['**Giai đoạn 1**\n(n ≥ 60 BN)', 'Đủ power cho so sánh 2 nhóm', 'Chạy Bảng A1 (biomarker theo giới), B1 (ROC), B4 (MDR vs nhạy)', 'A1, B1, B4 — kiểm tra sơ bộ ngưỡng cắt'],
        ['**Giai đoạn 2**\n(n ≥ 100 BN)', 'Đủ power cho logistic regression', 'Bảng A2 (pattern cytokine), B2 (mô hình tổ hợp), B3 (TIMP-1 theo giới), A3 (đồng nhiễm)', 'Hồi quy đa biến; DeLong test; ROC phân tầng giới'],
        ['**Tùy chọn**\n(Nếu có dữ liệu dọc)', 'Lấy mẫu ngày 3–5 và ra viện', 'Biomarker kinetics — xây dựng protocol bổ sung', 'Survival analysis; repeated measures; bài báo độc lập'],
    ],
    col_widths=[3.5, 3, 5, 5]
)

doc.add_paragraph()
add_heading('7.2. Đánh giá theo tiêu chí ưu tiên', level=2)
add_table(
    headers=['Đề xuất phân tích', 'Tính mới\n(Việt Nam)', 'Khả thi\n(dữ liệu có sẵn)', 'Giá trị\nthực hành', 'Ưu tiên'],
    rows=[
        ['TIMP-1 phân tầng theo giới (B3)', '★★★★★', '★★★★', '★★★', '⭐ Cao nhất'],
        ['Pattern cytokine theo VK (A2)', '★★★★★', '★★★ (cần PCR)', '★★★★', '⭐ Cao'],
        ['Mô hình tổ hợp PSI + biomarker (B2)', '★★★★', '★★★★', '★★★★', '⭐ Cao'],
        ['Ngưỡng cắt NLR/PLR/CAR cho VN (B1)', '★★★', '★★★★★', '★★★★★', '⭐ Cao'],
        ['Biomarker vs MDR/XDR (B4)', '★★★★', '★★★★', '★★★★★', '⭐ Cao'],
        ['Tỷ lệ đồng nhiễm – mức độ nặng (A3)', '★★★', '★★★ (cần PCR)', '★★★★', 'Trung bình'],
        ['Biomarker kinetics – dọc', '★★★★★', '★★ (cần protocol)', '★★★★★', 'Tùy chọn'],
    ],
    col_widths=[6, 2.5, 2.5, 2.5, 3]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  PHẦN VIII — TÀI LIỆU THAM KHẢO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('VIII. TÀI LIỆU THAM KHẢO', level=1)
add_hr()
add_info_box('Nguồn tra cứu: Consensus MCP, PubMed Extended MCP, Scopus MCP, Scholar-Sidekick MCP, YHVN BigQuery MCP, Crossref MCP, Clinical Trials MCP, WebSearch. Format: Vancouver. Kiểm tra retraction: Scholar-Sidekick (Crossref/Retraction Watch) — 49/49 tài liệu không bị thu hồi.')

refs = [
    '[1] Aladakatti AR et al. Soluble sTREM-1 as a Prognostic Biomarker of Mortality in Ventilator-Associated Pneumonia: A Prospective Observational Study. Respir Med. 2025.',
    '[2] Hogendoorn SKL et al. Clinical sign and biomarker-based algorithm to identify bacterial pneumonia among outpatients with lower respiratory tract infection in Tanzania. BMC Infect Dis. 2022;22:39.',
    '[3] Tekin A et al. The Neutrophil/Lymphocyte Ratio and Outcomes in Hospitalized Patients with Community-Acquired Pneumonia: A Retrospective Cohort Study. Biomedicines. 2024.',
    '[4] Menéndez R et al. Cytokine Activation Patterns and Biomarkers Are Influenced by Microorganisms in Community-Acquired Pneumonia. Chest. 2012;141(6):1537–1545.',
    '[5] Wang Y et al. The usefulness of serum PCT, CRP, sTREM-1 and CPIS for evaluation of severity and prognosis of CAP in elderly patients. Arch Gerontol Geriatr. 2019;80:53–57.',
    '[6] Tejera A et al. Prognosis of CAP: value of TREM-1 and other mediators of the inflammatory response. Cytokine. 2007;38(3):117–123. PMID: 17659879.',
    '[7] How CK et al. Usefulness of TREM-1 in differentiating between typical and atypical CAP. Am J Emerg Med. 2011;29(5):531–536. PMID: 21114709.',
    '[8] Almuntashiri S et al. Plasma TIMP-1 as a sex-specific biomarker for acute lung injury. Biol Sex Differ. 2022;13(1):70. PMID: 36482481.',
    '[9] Almuntashiri S et al. TIMP-1 and its potential diagnostic and prognostic value in pulmonary diseases. Chin Med J Pulm Crit Care Med. 2023;1(2):67–76.',
    '[10] Almuntashiri S et al. Estrogen-dependent gene regulation: Molecular basis of TIMP-1 as a sex-specific biomarker for ALI. Physiol Rep. 2024;12. DOI: 10.14814/phy2.70047.',
    '[11] Jones T et al. Elevated Plasma Levels of MMP-3 and TIMP-1 Associate With Organ Dysfunction and Mortality in Sepsis. Shock. 2021.',
    '[12] Lorente L et al. MMP-9, -10, and TIMP-1 blood levels as biomarkers of severity and mortality in sepsis. Crit Care. 2009;13(5):R178.',
    '[13] Feng CM et al. Serum interleukin-17 predicts severity and prognosis in patients with CAP: a prospective cohort study. BMC Pulm Med. 2021;21. PMID: 34856971.',
    '[14] Dao BN et al. Relationship between serum TNF-α, IL-6, and IL-10 levels and disease severity in bacterial CAP. Pneumon. 2023;36(4):1–8.',
    '[15] Miyazaki T et al. A high α1-antitrypsin/interleukin-10 ratio predicts bacterial pneumonia in adults with CAP. Pneumonia. 2023.',
    '[16] Ganaie ZA et al. Association Between Elevated NLR and Mortality Risk in CAP: A Systematic Review and Meta-Analysis. Cureus. 2025. PMID: 41146762.',
    '[17] Kuikel S et al. Neutrophil-lymphocyte ratio as a predictor of adverse outcome in CAP: A systematic review. Health Sci Rep. 2022.',
    '[18] Huang L et al. The improved prediction value of NLR to pneumonia severity scores for mortality in older people with CAP. BMC Geriatrics. 2025.',
    '[19] Ustaalioğlu İ et al. CRP-to-albumin ratio as a prognostic marker in CAP mortality. Turk J Clinics Lab. 2025.',
    '[20] Luo B et al. Two new inflammatory markers related to the CURB-65 score for disease severity in CAP. Open Life Sci. 2021;16(1):84–91.',
    '[21] Wang M et al. Predicting severe CAP in adults: a machine learning approach using CAR. J Thorac Dis. 2025.',
    '[22] Viasus D et al. Biomarkers for predicting short-term mortality in CAP: A systematic review and meta-analysis. J Infect. 2016.',
    '[23] Çetin E et al. Advancing ICU mortality prediction in CAP: Combining FAR, CT severity score, PSI, and CURB-65. Biomol Biomed. 2025.',
    '[24] Menéndez R et al. Biomarkers improve mortality prediction by prognostic scales in CAP. Thorax. 2009.',
    '[25] Tran H et al. Community-acquired pneumonia-causing bacteria and antibiotic resistance rate among Vietnamese patients: A cross-sectional study. Medicine. 2022. PMID: 36086715.',
    '[26] Tran Thi Ngoc Dung et al. The bacterial etiology and antimicrobial susceptibility of LRTI in Vietnam. Ann Clin Microbiol Antimicrob. 2025.',
    '[27] Phan Hồng Thảo Vy et al. Phát hiện tác nhân vi sinh bằng Multiplex-time PCR ở bệnh nhân CAP nặng tại BV ĐK TW Cần Thơ. Tạp chí Y học Việt Nam. 2025;552(3):294–298.',
    '[28] Lý Khánh Vân et al. Tác nhân vi sinh gây CAP ở người lớn nhập viện có hay không có đái tháo đường. Tạp chí Y học Việt Nam. 2025;554(1):353–357.',
    '[29] Nguyễn Thanh Huyền et al. Xác định căn nguyên vi rút gây CAP bằng Realtime PCR đa mồi tại BV Bạch Mai. Tạp chí Y học Việt Nam. 2025;551(3):134–138.',
    '[30] Corica B et al. Sex and gender differences in community-acquired pneumonia. Intern Emerg Med. 2022;17:1279–1292.',
    '--- Bổ sung Session 2 (2026-06-11) ---',
    '[31] Cataudella E et al. Neutrophil-to-lymphocyte ratio: an emerging marker predicting prognosis in elderly adults with community-acquired pneumonia. J Am Geriatr Soc. 2017;65(8):1796–1801. PMID: 28543831.',
    '[32] Baran B et al. Assessment of mortality risk in patients with community-acquired pneumonia: role of novel inflammatory biomarkers (CALLY index). J Clin Lab Anal. 2025.',
    '[33] Moravec M et al. Th17 lymphocytes and IL-17A during the course of severe community-acquired pneumonia, comparison with etiology and outcome. Epidemiol Mikrobiol Imunol. 2024.',
    '[34] Karamouzos V et al. Cytokine production and outcome in MDR versus non-MDR gram-negative bacteraemia and sepsis. Infect Dis. 2021;53(3):199–208.',
    '[35] Wang H et al. High mortality of Acinetobacter baumannii infection is attributed to macrophage-mediated induction of cytokine storm but preventable by naproxen. eBioMedicine. 2024. PMID: 38518726.',
    '[36] Kumar NR et al. Multidrug-resistant sepsis: a critical healthcare challenge. Antibiotics. 2024;13(2):119.',
    '[37] Bạch Thái Dương, Nguyễn Quang Thái, Nguyễn Thị Ngọc Lan và cộng sự. Vi khuẩn đa kháng thuốc trên bệnh nhân viêm phổi mắc phải cộng đồng ở người lớn tại Bệnh viện Đa khoa Trung ương Cần Thơ. Tạp chí Y học Việt Nam. 2024;543(3). DOI: 10.51298/vmj.v543i3.11580.',
    '[38] Lê Thị Diệu Hiền, Mai Xuân Khẩn, Tạ Bá Thắng. Thay đổi nồng độ cytokine huyết thanh ở bệnh nhân viêm phổi cộng đồng do vi khuẩn. Tạp chí Y học Việt Nam. 2021;502(2). DOI: 10.51298/vmj.v502i2.642.',
    '[39] Đỗ Thanh Hoà, Lê Đức Giang. So sánh giá trị của một số thang điểm trong dự đoán nhập đơn vị hồi sức tích cực ở bệnh nhân viêm phổi mắc phải tại cộng đồng. Tạp chí Y học Việt Nam. 2025;557(1). DOI: 10.51298/vmj.v557i1.16612.',
    '[40] Duan Y et al. Serum cytokine levels in children with community-acquired pneumonia caused by different respiratory pathogens. Ital J Pediatr. 2025;51. DOI: 10.1186/s13052-025-01901-7.',
    '[41] Esposito S et al. Sensitivity and specificity of soluble triggering receptor expressed on myeloid cells-1, midregional proatrial natriuretic peptide and midregional proadrenomedullin for distinguishing etiology and to assess severity in community-acquired pneumonia. PLoS ONE. 2016;11(9):e0163262. DOI: 10.1371/journal.pone.0163262.',
    '[42] Mou S et al. Serum level of soluble triggering receptor expressed on myeloid cells-1 in patients with community-acquired pneumonia complicated with COPD. Can Respir J. 2022. DOI: 10.1155/2022/4494756.',
    '[43] Lê Thị Huệ, Lý Khánh Vân, Hoàng Tiến Mỹ, Phan Thị Cẩm Luyến. Phối hợp tác nhân vi khuẩn, virus trên bệnh nhân viêm phổi mắc phải cộng đồng nhập viện. Tạp chí Y học Việt Nam. 2023;526(2). DOI: 10.51298/vmj.v526i2.5568.',
    '[44] Lý Khánh Vân, Lê Thị Huệ, Phan Thị Cẩm Luyến và cộng sự. Tác nhân vi sinh phát hiện trên mẫu đàm bệnh nhân viêm phổi mắc phải cộng đồng nhập viện. Tạp chí Y học Việt Nam. 2023;530(1). DOI: 10.51298/vmj.v530i1.6597.',
    '[45] Lý Khánh Vân, Phạm Hùng Vân. Các gene sinh carbapenemase và sự đề kháng kháng sinh của Acinetobacter baumannii ở bệnh nhân viêm phổi mắc phải cộng đồng nhập viện. Tạp chí Y học Việt Nam. 2025;548(1). DOI: 10.1055/s-0037-1599225.',
    '[46] Phan Hồng Thảo Vy, Đỗ Hoàng Long, Lê Minh Nhân. Phát hiện tác nhân vi sinh bằng multiplex real-time PCR ở bệnh nhân viêm phổi cộng đồng nặng tại Bệnh viện Đa khoa Trung ương Cần Thơ. Tạp chí Y học Việt Nam. 2025;552(3). DOI: 10.51298/vmj.v552i2.15140.',
    '[47] Jiang Y et al. Neutrophil-to-lymphocyte ratio in Mycoplasma pneumoniae pneumonia: a systematic review and meta-analysis. BMC Pulm Med. 2026. DOI: 10.1186/s12890-026-03000-x.',
    '[48] Zhang X et al. Interleukins in community-acquired pneumonia: from biomarkers to precision medicine. Front Immunol. 2026;17:1774731. DOI: 10.3389/fimmu.2026.1774731.',
    '[49] Nguyễn Thị Pháp, Phan Vũ Nguyên. So sánh giá trị tiên lượng tử vong 30 ngày theo tiêu chuẩn phụ viêm phổi nặng ATS/IDSA, thang điểm CURB-65 và PSI ở bệnh nhân viêm phổi mắc phải cộng đồng tại Bệnh viện Phạm Ngọc Thạch. Tạp chí Y học Việt Nam. 2024;545(1). DOI: 10.51298/vmj.v545i1.12138.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.style = 'Normal'
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  PHỤ LỤC — BẢNG TÓM TẮT BIOMARKER
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('PHỤ LỤC: BẢNG TÓM TẮT BIOMARKER — HIỆU QUẢ DỰ ĐOÁN', level=1)
add_hr()
add_table(
    headers=['Biomarker', 'Nghiên cứu tốt nhất', 'AUC tốt nhất', 'Cutoff', 'Ứng dụng chính', 'Điểm mới tại VN'],
    rows=[
        ['sTREM-1', 'Aladakatti 2025\n+ APACHE II', '0,945 (tổ hợp)\n0,83 (đơn độc)', '>9,5–44,2 pg/mL', 'Phân biệt VK điển hình vs KĐH;\nTiên lượng tử vong', '★★★★★\nKhông có dữ liệu VN'],
        ['**TIMP-1**\n(SEX-SPECIFIC)', 'Almuntashiri 2022\n(ARDS, nữ)', '0,87 (nữ)\nKém ý nghĩa (nam)', '159,7 ng/mL\n(chỉ cho nữ)', 'Tử vong 30 và 90 ngày\nĐặc hiệu giới tính', '★★★★★\nHoàn toàn mới tại châu Á'],
        ['IL-17', 'Feng et al. 2021\n(CAP người lớn)', '0,89 (tử vong)', '86,80 ng/mL', 'Tiên lượng tử vong\nNhập ICU, thở máy', '★★★★\nChưa có VN người lớn'],
        ['IL-6', 'Menéndez 2009\n(PSI + CRP + IL-6)', 'PSI: 0,80 → 0,85\n(thêm CRP)', '—', 'Tăng AUC PSI\nPhân biệt S. pneumoniae', '★★★\nDao 2023 có 1 NC VN'],
        ['IL-10', 'Miyazaki 2023\n(AAT/IL-10 >65)', 'AUC panel 0,927', 'Tỷ số AAT/IL-10 >65', 'Phân biệt VK vs virus', '★★★★\nChưa có VN'],
        ['NLR', 'Ganaie 2025\n(meta, 17.838 BN)', 'Pooled RR 2,02\nAUC ~0,72–0,76', '>10 (phổ biến)\n>6,5 (người già)', 'ICU admission\nTử vong (sau PSI)', '★★★\nNgưỡng VN chưa xác định'],
        ['CAR', 'Ustaalioğlu 2025\nn=312 CAP', '0,837', '>0,77', 'Tử vong 30 ngày\nCAP nặng (OR 8,8)', '★★★\nNgưỡng VN chưa xác định'],
        ['PLR', 'Enersen 2023\n(3.826 BN đa trung tâm)', 'HR 1,001 (90 ngày)', '—', 'Tử vong 90 ngày\n(yếu hơn NLR)', '★★\nNgưỡng VN chưa xác định'],
    ],
    col_widths=[2.5, 3, 2.5, 2.5, 3.5, 3]
)

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Hết báo cáo —\nTháng 6 năm 2026 | Đề tài Tiến sĩ CAP | Bệnh viện Phổi Hải Phòng')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = RGBColor(127, 127, 127)

# SAVE
out = '/Users/buiminhkhoi/Documents/Antigravity/cap-research/outputs/cap_analysis_report.docx'
doc.save(out)
print(f'✓ Đã lưu: {out}')
